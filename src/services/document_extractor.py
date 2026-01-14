"""
Document Extractor - extracts text from Word documents (.docx, .doc).
"""
import logging
from pathlib import Path

from docx import Document

from src.config import MAX_FILE_SIZE, MAX_DOCUMENT_PAGES
from src.services.extracted_content import ExtractedContent

logger = logging.getLogger(__name__)

# Approximate characters per page for page count estimation
CHARS_PER_PAGE = 2000


class DocumentExtractor:
    """Extracts text from Word documents."""

    async def extract(
        self,
        file_path: str | Path,
    ) -> ExtractedContent:
        """
        Extract text from Word document.

        Args:
            file_path: Path to .docx file

        Returns:
            ExtractedContent with extracted text
        """
        file_path = Path(file_path)

        # Check file exists
        if not file_path.exists():
            return ExtractedContent.from_error(
                "Файл не найден", source_type="docx"
            )

        # Check file extension
        suffix = file_path.suffix.lower()
        if suffix not in ('.docx', '.doc'):
            return ExtractedContent.from_error(
                f"Формат {suffix} не поддерживается. Используйте .docx",
                source_type="docx"
            )

        # Check file size
        file_size = file_path.stat().st_size
        if file_size > MAX_FILE_SIZE:
            return ExtractedContent.from_error(
                f"Документ слишком большой ({file_size // 1024 // 1024}MB). "
                f"Максимум: {MAX_FILE_SIZE // 1024 // 1024}MB",
                source_type="docx"
            )

        # .doc files are not supported by python-docx
        if suffix == '.doc':
            return ExtractedContent.from_error(
                "Старый формат .doc не поддерживается. "
                "Пожалуйста, сохраните документ в формате .docx",
                source_type="docx"
            )

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        paragraphs.append(" | ".join(row_texts))

            if not paragraphs:
                return ExtractedContent.from_error(
                    "Документ пуст или не содержит текста",
                    source_type="docx"
                )

            full_text = "\n\n".join(paragraphs)

            # Estimate page count
            estimated_pages = max(1, len(full_text) // CHARS_PER_PAGE)

            if estimated_pages > MAX_DOCUMENT_PAGES:
                return ExtractedContent.from_error(
                    f"Документ слишком длинный (~{estimated_pages} страниц). "
                    f"Максимум: {MAX_DOCUMENT_PAGES}",
                    source_type="docx"
                )

            # Try to extract title from first paragraph or document properties
            title = self._extract_title(doc, paragraphs)

            return ExtractedContent(
                text=full_text,
                title=title,
                source_type="docx",
                metadata={
                    "file_size": file_size,
                    "paragraph_count": len(paragraphs),
                    "estimated_pages": estimated_pages,
                    "table_count": len(doc.tables),
                }
            )

        except Exception as e:
            logger.error(f"Document extraction error: {e}")
            return ExtractedContent.from_error(
                f"Ошибка извлечения текста из документа: {str(e)}",
                source_type="docx"
            )

    def _extract_title(self, doc: Document, paragraphs: list[str]) -> str | None:
        """Try to extract title from document."""
        # Try document core properties
        try:
            if doc.core_properties.title:
                return doc.core_properties.title
        except Exception:
            pass

        # Try first heading or paragraph
        if paragraphs:
            first = paragraphs[0]
            # Check if it looks like a title (not too long, not too short)
            if 5 < len(first) < 150:
                return first

        return None
